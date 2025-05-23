#!/usr/bin/env python3
# Main application file for HFE mutation analysis

import pandas as pd
import numpy as np
import os
import tempfile
from openpyxl.styles import Font
from pathlib import Path
import io
import base64
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import Shiny for Python
from shiny import App, ui, render, reactive

# Import modules
from modules.data_processing import clean_dataset, analyze_age_column, find_hfe_columns, clean_column_names
from modules.analysis import (
    analyze_dataset, analyze_hh_risk, analyze_diagnosis_associations, 
    analyze_hardy_weinberg
)
from modules.visualization import (
    generate_genotype_distribution_plot, generate_risk_distribution_plot,
    generate_diagnosis_association_plot, generate_liver_disease_plot,
    generate_hardy_weinberg_plots, generate_genotype_by_age_plot,
    generate_genotype_by_gender_plot, generate_genotype_diagnosis_gender_plot
)
from modules.ui import create_ui
from modules.diagnosis_trends import analyze_diagnosis_trends, generate_diagnosis_validation_report

# Define the Shiny server
def server(input, output, session):
    # Reactive value to store the dataset
    data = reactive.Value(None)
    data_cleaned = reactive.Value(False)
    analysis_run = reactive.Value(False)
    analysis_results_text = reactive.Value([])
    data_info = reactive.Value([])
    plots_genotype = reactive.Value([])
    plot_risk = reactive.Value(None)
    plot_diagnosis = reactive.Value(None)
    plot_liver = reactive.Value(None)
    plots_hardy_weinberg = reactive.Value([])
    
    # Reactive values for the visualizations
    plots_age = reactive.Value([])
    plots_gender = reactive.Value([])
    plots_diagnosis_gender = reactive.Value([])
    
    # Diagnostic reactive values
    diagnostics_age = reactive.Value([])
    diagnostics_gender = reactive.Value([])
    diagnostics_diagnosis_gender = reactive.Value([])
    
    # New reactive values for diagnosis trends
    trends_results_data = reactive.Value([])
    trends_plots_data = reactive.Value([])
    validation_results_data = reactive.Value([])
    
    # Add JavaScript to toggle diagnosis trends options visibility
    ui.insert_ui(
        ui.tags.script("""
        $(document).ready(function() {
            function updateVisibility() {
                var selectedTypes = $('input[name="analysis_type"]:checked').map(function() {
                    return $(this).val();
                }).get();
                
                // Show diagnosis trends options if diagnosis_trends is selected
                if (selectedTypes.includes('diagnosis_trends')) {
                    $('#diagnosis_trends_options').show();
                } else {
                    $('#diagnosis_trends_options').hide();
                }
            }
            
            $('input[name="analysis_type"]').on('change', updateVisibility);
            updateVisibility(); // Initial update
        });
        """),
        selector="head",
        immediate=True
    )
    
    @reactive.Effect
    @reactive.event(input.file1)
    def _():
        if input.file1() is None:
            return
        
        file_info = input.file1()
        file_path = file_info[0]["datapath"]
        
        try:
            df = pd.read_excel(file_path)
            
            # Clean column names immediately after loading
            original_cols = list(df.columns)
            df = clean_column_names(df)
            cleaned_cols = list(df.columns)
            
            data.set(df)
            data_cleaned.set(False)
            analysis_run.set(False)
            
            # Basic dataset info
            info = []
            info.append(f"File: {file_info[0]['name']}")
            info.append(f"Rows: {len(df)}")
            info.append(f"Columns: {len(df.columns)}")
            
            # Report on column name cleaning if any were changed
            cleaned_count = sum(1 for i, col in enumerate(original_cols) if col != cleaned_cols[i])
            if cleaned_count > 0:
                info.append(f"\nCleaned {cleaned_count} column names by removing trailing spaces:")
                for i, (orig, cleaned) in enumerate(zip(original_cols, cleaned_cols)):
                    if orig != cleaned:
                        info.append(f"  Column {i}: '{orig}' → '{cleaned}'")
            
            # Show column names to help user select date column
            info.append("\nAvailable columns:")
            for i, col in enumerate(df.columns):
                info.append(f"  {i}: {col}")
            
            # HFE columns
            hfe_columns = find_hfe_columns(df)
            
            if hfe_columns:
                info.append(f"\nFound {len(hfe_columns)} HFE mutation columns:")
                for col in hfe_columns:
                    info.append(f"- {col}")
            else:
                info.append("\nNo HFE mutation columns found in the dataset")
            
            data_info.set(info)
        except Exception as e:
            ui.notification_show(f"Error loading file: {str(e)}", type="error", duration=None)
    
    @reactive.Effect
    @reactive.event(input.btn_clean)
    def _():
        if data() is None:
            ui.notification_show("Please upload a dataset first", type="warning")
            return
        
        try:
            df, results = clean_dataset(
                data(), 
                remove_blank_validovany=input.clean_validovany(),
                remove_blank_diagnoza=input.clean_diagnoza(),
                remove_blank_hfe=input.clean_hfe(),
                remove_second_column=input.remove_second_col(),
                min_age=input.min_age(),
                max_age=input.max_age(),
                filter_by_age=input.filter_by_age()
            )
            data.set(df)
            data_cleaned.set(True)
            analysis_run.set(False)
            data_info.set(results)
            ui.notification_show("Dataset cleaned successfully", type="success")
        except Exception as e:
            ui.notification_show(f"Error cleaning dataset: {str(e)}", type="error", duration=None)
    
    @reactive.Effect
    @reactive.event(input.btn_analyze)
    def _():
        if data() is None:
            ui.notification_show("Please upload a dataset first", type="warning")
            return
        
        if not data_cleaned():
            ui.notification_show("It's recommended to clean the dataset before analysis", type="info")
        
        try:
            df = data()
            results = []
            
            # Find HFE columns
            hfe_columns = find_hfe_columns(df)
            
            if not hfe_columns:
                ui.notification_show("No HFE mutation columns found in the dataset", type="warning")
                return
            
            # Add information about dataset
            results.append("Dataset columns:")
            for i, col in enumerate(df.columns):
                results.append(f"  {i}: {col}")
            
            # Get selected analysis types
            selected_analysis = input.analysis_type()
            
            # Run selected analyses
            if not selected_analysis:
                ui.notification_show("Please select at least one analysis type", type="warning")
                return
            
            if "basic" in selected_analysis:
                basic_results = analyze_dataset(df)
                results.extend(basic_results)
                
                # Add specific age column analysis
                age_results = analyze_age_column(df)
                results.extend(age_results)
            
            if "hh_risk" in selected_analysis:
                df, hh_results = analyze_hh_risk(df, hfe_columns)
                results.extend(hh_results)
            
            if "diagnosis" in selected_analysis:
                df, diag_results = analyze_diagnosis_associations(df, hfe_columns)
                results.extend(diag_results)
            
            if "hardy_weinberg" in selected_analysis:
                # Get selected Hardy-Weinberg tests
                selected_hwe_tests = input.hwe_tests() if input.hwe_tests() else ["chi_square"]
                
                # Run Hardy-Weinberg analysis with selected tests
                df, hwe_results, hwe_results_list = analyze_hardy_weinberg(df, hfe_columns, selected_hwe_tests)
                results.extend(hwe_results)
                plots_hardy_weinberg.set(generate_hardy_weinberg_plots(hwe_results_list))
            
            if "diagnosis_trends" in selected_analysis:
                # Get column name from user input if auto-detect is disabled
                selected_date_col = input.date_column() if not input.auto_detect_date() else None
                
                # Run analysis with selected column name
                trends_results, trends_plots = analyze_diagnosis_trends(df, selected_date_col)
                trends_results_data.set(trends_results)
                trends_plots_data.set(trends_plots)
                
                # Find diagnosis column for validation report
                diagnoza_col = next((col for col in df.columns 
                                   if any(term in str(col).lower() for term in 
                                         ['diagnoza', 'mkch', 'diagnosis', 'icd', 'kod'])), None)
                
                if diagnoza_col:
                    # Generate validation report
                    validation_results, _ = generate_diagnosis_validation_report(df, diagnoza_col)
                    validation_results_data.set(validation_results)
            
            # Always generate genotype distribution plots
            plots_genotype.set(generate_genotype_distribution_plot(df, hfe_columns))
            
            # Generate new plots for age, gender, and diagnosis-gender relationships
            age_plots, age_diagnostics = generate_genotype_by_age_plot(df, hfe_columns)
            plots_age.set(age_plots)
            diagnostics_age.set(age_diagnostics)
            
            gender_plots, gender_diagnostics = generate_genotype_by_gender_plot(df, hfe_columns)
            plots_gender.set(gender_plots)
            diagnostics_gender.set(gender_diagnostics)
            
            diag_gender_plots, diag_gender_diagnostics = generate_genotype_diagnosis_gender_plot(df, hfe_columns)
            plots_diagnosis_gender.set(diag_gender_plots)
            diagnostics_diagnosis_gender.set(diag_gender_diagnostics)
            
            if "hh_risk" in selected_analysis:
                plot_risk.set(generate_risk_distribution_plot(df))
            
            if "diagnosis" in selected_analysis:
                plot_diagnosis.set(generate_diagnosis_association_plot(df))
                plot_liver.set(generate_liver_disease_plot(df))
            
            # Update data with analysis results
            data.set(df)
            analysis_results_text.set(results)
            analysis_run.set(True)
            ui.notification_show("Analysis completed", type="success")
        
        except Exception as e:
            ui.notification_show(f"Error during analysis: {str(e)}", type="error", duration=None)
            import traceback
            error_traceback = traceback.format_exc()
            analysis_results_text.set([f"Error during analysis: {str(e)}", "", "Traceback:", error_traceback])
            analysis_run.set(True)
    
    @output
    @render.text
    def dataset_info():
        if data() is None:
            return "No dataset loaded"
        return "\n".join(data_info())
    
    @output
    @render.data_frame
    def data_preview():
        if data() is None:
            return None
        return render.DataGrid(data().head(10), width="100%")
    
    @output
    @render.text
    def analysis_results():
        if not analysis_run():
            return "Run analysis to see results"
        return "\n".join(analysis_results_text())
    
    @output
    @render.ui
    def genotype_plots():
        if not plots_genotype():
            return ui.p("No genotype plots available. Run analysis first.")
        
        plot_htmls = []
        for name, img_str in plots_genotype():
            plot_htmls.append(ui.tags.div(
                ui.tags.h4(name),
                ui.tags.img(src=f"data:image/png;base64,{img_str}", width="100%", style="max-width: 800px;"),
                ui.br(), ui.br()
            ))
        
        return ui.tags.div(*plot_htmls)
    
    @output
    @render.ui
    def genotype_age_plots():
        if not plots_age() or len(plots_age()) == 0:
            diagnostics = diagnostics_age()
            if diagnostics:
                # Show the most relevant diagnostic information
                key_messages = [msg for msg in diagnostics if any(term in msg for term in 
                              ['column', 'found', 'Sample', 'Age range', 'Error'])]
                
                # Limit to 5 most important messages
                if len(key_messages) > 5:
                    key_messages = key_messages[:5]
                    
                return ui.tags.div(
                    ui.p("No age plots available. Diagnostic information:"),
                    ui.tags.ul(*[ui.tags.li(msg) for msg in key_messages]),
                    ui.p("See Analysis Results tab for complete diagnostics.")
                )
            else:
                return ui.p("No age plots available. This could be because age data is missing or there aren't enough samples to create meaningful visualizations.")
        
        plot_htmls = []
        for name, img_str in plots_age():
            plot_htmls.append(ui.tags.div(
                ui.tags.h4(name),
                ui.tags.img(src=f"data:image/png;base64,{img_str}", width="100%", style="max-width: 800px;"),
                ui.br(), ui.br()
            ))
        
        return ui.tags.div(*plot_htmls)
    
    @output
    @render.ui
    def genotype_gender_plots():
        if not plots_gender() or len(plots_gender()) == 0:
            diagnostics = diagnostics_gender()
            if diagnostics:
                # Show the most relevant diagnostic information
                key_messages = [msg for msg in diagnostics if any(term in msg for term in 
                              ['column', 'found', 'Sample', 'distribution', 'Error'])]
                
                # Limit to 5 most important messages
                if len(key_messages) > 5:
                    key_messages = key_messages[:5]
                    
                return ui.tags.div(
                    ui.p("No gender plots available. Diagnostic information:"),
                    ui.tags.ul(*[ui.tags.li(msg) for msg in key_messages]),
                    ui.p("See Analysis Results tab for complete diagnostics.")
                )
            else:
                return ui.p("No gender plots available. This could be because gender data is missing or there aren't enough samples to create meaningful visualizations.")
        
        plot_htmls = []
        for name, img_str in plots_gender():
            plot_htmls.append(ui.tags.div(
                ui.tags.h4(name),
                ui.tags.img(src=f"data:image/png;base64,{img_str}", width="100%", style="max-width: 800px;"),
                ui.br(), ui.br()
            ))
        
        return ui.tags.div(*plot_htmls)
    
    @output
    @render.ui
    def genotype_diagnosis_gender_plots():
        if not plots_diagnosis_gender() or len(plots_diagnosis_gender()) == 0:
            diagnostics = diagnostics_diagnosis_gender()
            if diagnostics:
                # Show the most relevant diagnostic information
                key_messages = [msg for msg in diagnostics if any(term in msg for term in 
                              ['column', 'found', 'Sample', 'distribution', 'Error'])]
                
                # Limit to 5 most important messages
                if len(key_messages) > 5:
                    key_messages = key_messages[:5]
                    
                return ui.tags.div(
                    ui.p("No diagnosis-gender plots available. Diagnostic information:"),
                    ui.tags.ul(*[ui.tags.li(msg) for msg in key_messages]),
                    ui.p("See Analysis Results tab for complete diagnostics.")
                )
            else:
                return ui.p("No diagnosis-gender plots available. This could be because diagnosis or gender data is missing or there aren't enough samples to create meaningful visualizations.")
        
        plot_htmls = []
        for name, img_str in plots_diagnosis_gender():
            plot_htmls.append(ui.tags.div(
                ui.tags.h4(name),
                ui.tags.img(src=f"data:image/png;base64,{img_str}", width="100%", style="max-width: 800px;"),
                ui.br(), ui.br()
            ))
        
        return ui.tags.div(*plot_htmls)
    
    @output
    @render.ui
    def risk_plot():
        if not plot_risk():
            return ui.p("No risk distribution plot available. Run HH Risk analysis first.")
        
        return ui.tags.div(
            ui.tags.img(src=f"data:image/png;base64,{plot_risk()}", width="100%", style="max-width: 800px;")
        )
    
    @output
    @render.ui
    def diagnosis_plot():
        if not plot_diagnosis():
            return ui.p("No diagnosis association plot available. Run Diagnosis Association analysis first.")
        
        return ui.tags.div(
            ui.tags.img(src=f"data:image/png;base64,{plot_diagnosis()}", width="100%", style="max-width: 800px;")
        )
    
    @output
    @render.ui
    def liver_disease_plot():
        if not plot_liver():
            return ui.p("No liver disease plot available. Run Diagnosis Association analysis first.")
        
        return ui.tags.div(
            ui.tags.img(src=f"data:image/png;base64,{plot_liver()}", width="100%", style="max-width: 800px;")
        )
    
    @output
    @render.ui
    def hardy_weinberg_plots():
        if not plots_hardy_weinberg():
            return ui.p("No Hardy-Weinberg plots available. Run Hardy-Weinberg analysis first.")
        
        plot_htmls = []
        for mutation, img_str in plots_hardy_weinberg():
            plot_htmls.append(ui.tags.div(
                ui.tags.h4(mutation),
                ui.tags.img(src=f"data:image/png;base64,{img_str}", width="100%", style="max-width: 800px;"),
                ui.br(), ui.br()
            ))
        
        return ui.tags.div(*plot_htmls)
    
    @output
    @render.text
    def diagnosis_trends_results():
        if not analysis_run() or not trends_results_data():
            return "Run Diagnosis Trends Analysis to see results"
        return "\n".join(trends_results_data())
    
    @output
    @render.text
    def diagnosis_validation_results():
        if not analysis_run() or not validation_results_data():
            return "Run Diagnosis Trends Analysis to see validation results"
        return "\n".join(validation_results_data())
    
    @output
    @render.ui
    def diagnosis_trends_plots():
        if not trends_plots_data():
            return ui.p("No diagnosis trends plots available. Run diagnosis trends analysis first.")
        
        plot_htmls = []
        for name, img_str in trends_plots_data():
            plot_htmls.append(ui.tags.div(
                ui.tags.h4(name),
                ui.tags.img(src=f"data:image/png;base64,{img_str}", width="100%", style="max-width: 800px;"),
                ui.br(), ui.br()
            ))
        
        return ui.tags.div(*plot_htmls)
    
    @reactive.Effect
    @reactive.event(data_cleaned)
    def _():
        # This effect will run whenever data_cleaned changes, but we'll manage
        # the button states through the disabled attribute in the UI
        pass
        
    @reactive.Effect
    @reactive.event(data)
    def _():
        # This effect will run whenever data changes, but we'll manage
        # the button states through the disabled attribute in the UI
        pass
        
    # Helper functions to determine button states
    @reactive.calc
    def formatted_button_disabled():
        return not (data() is not None and data_cleaned())
    
    @reactive.calc
    def download_button_disabled():
        return data() is None
    
    @output
    @render.ui
    def formatted_download_button():
        return ui.download_button(
            "download_formatted", 
            "Save Formatted Dataset", 
            class_="btn-outline-secondary mt-2", 
            disabled=formatted_button_disabled()
        )
    
    @output
    @render.ui
    def data_download_button():
        return ui.download_button(
            "download_data", 
            "Download Processed Data", 
            class_="btn-info", 
            disabled=download_button_disabled()
        )
    
    @output
    @render.download(filename="processed_dataset.xlsx")
    def download_data():
        # If no data, return None
        if data() is None:
            return None
            
        # Create temporary file name
        temp_path = os.path.join(tempfile.gettempdir(), "processed_dataset.xlsx")
        
        # Get data and format ID column
        df = data().copy()
        
        # Find ID column (likely first column)
        if 'id' in df.columns:
            id_col = 'id'
        elif any(col.lower() == 'id' for col in df.columns):
            id_col = next(col for col in df.columns if col.lower() == 'id')
        else:
            # Assuming the first column is the ID
            id_col = df.columns[0]
            
        # Fix ID formatting if needed
        if df[id_col].dtype in [np.int64, np.float64]:
            df[id_col] = df[id_col].astype('Int64')
            df[id_col] = df[id_col].astype(str)
            df[id_col] = df[id_col].replace(['nan', '<NA>'], None)
        
        # Apply zfill to preserve leading zeros
        mask = df[id_col].notna() & (df[id_col] != '')
        if any(mask):
            # Check for ID format
            sample_ids = df[id_col].dropna().head(10).tolist()
            id_lengths = [len(str(id_val)) for id_val in sample_ids if str(id_val) != 'nan' and str(id_val) != '']
            zfill_length = 10 if not id_lengths or max(id_lengths) > 9 else 9
            # Apply consistent formatting
            df.loc[mask, id_col] = df.loc[mask, id_col].str.zfill(zfill_length)
        
        # Write to Excel with formatting for ID column
        with pd.ExcelWriter(temp_path, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
            
            # Apply text format to ID column to preserve leading zeros
            worksheet = writer.sheets['Sheet1']
            id_col_index = list(df.columns).index(id_col) + 1  # +1 because Excel is 1-indexed
            for cell in worksheet.iter_cols(min_col=id_col_index, max_col=id_col_index, min_row=2):
                for x in cell:
                    x.number_format = '@'
        
        # Return the path as a string
        return temp_path
    
    @output
    @render.download(filename="formatted_dataset.xlsx")
    def download_formatted():
        # If no data, return None
        if data() is None:
            return None
            
        # Create temporary file name
        temp_path = os.path.join(tempfile.gettempdir(), "formatted_dataset.xlsx")
        
        # Get data and format ID column
        df = data().copy()
        
        # Find ID column (likely first column)
        if 'id' in df.columns:
            id_col = 'id'
        elif any(col.lower() == 'id' for col in df.columns):
            id_col = next(col for col in df.columns if col.lower() == 'id')
        else:
            # Assuming the first column is the ID
            id_col = df.columns[0]
            
        # Fix ID formatting - convert to string
        if df[id_col].dtype in [np.int64, np.float64]:
            # For numeric columns, only convert non-NaN values to string
            df[id_col] = df[id_col].astype('Int64')  # nullable integer type
            df[id_col] = df[id_col].astype(str)
            # Replace 'nan' or '<NA>' strings with None
            df[id_col] = df[id_col].replace(['nan', '<NA>'], None)
        
        # Apply zfill to non-null values to preserve leading zeros
        mask = df[id_col].notna() & (df[id_col] != '')
        if any(mask):
            # Determine ID format based on sample
            sample_ids = df[id_col].dropna().head(10).tolist()
            id_lengths = [len(str(id_val)) for id_val in sample_ids if str(id_val) != 'nan' and str(id_val) != '']
            zfill_length = 10 if not id_lengths or max(id_lengths) > 9 else 9
            # Apply zfill for consistent length
            df.loc[mask, id_col] = df.loc[mask, id_col].str.zfill(zfill_length)
        
        # Write to Excel with formatting
        with pd.ExcelWriter(temp_path, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
            
            # Apply formatting to Excel worksheet
            worksheet = writer.sheets['Sheet1']
            
            # Create a bold font object directly
            bold_font = Font(bold=True)
            
            # Format headers in bold
            for col_num, column_title in enumerate(df.columns, 1):
                cell = worksheet.cell(row=1, column=col_num)
                cell.font = bold_font
            
            # Set ID column to text format (to preserve leading zeros)
            id_col_index = list(df.columns).index(id_col) + 1  # +1 because Excel is 1-indexed
            for cell in worksheet.iter_cols(min_col=id_col_index, max_col=id_col_index, min_row=2):
                for x in cell:
                    x.number_format = '@'
            
            # Auto-adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        # Return the path as a string
        return temp_path

    def generate_word_report(analysis_results, df, plots=None):
        """
        Generate a Word document containing the analysis results and plots.
        
        Args:
            analysis_results: List of text results
            df: DataFrame with the analyzed data
            plots: Optional dictionary of plots to include
        
        Returns:
            BytesIO object containing the Word document
        """
        # Create a new Word document
        doc = docx.Document()
        
        # Set the title
        title = doc.add_heading('HFE Mutation Analysis Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a timestamp
        import datetime
        timestamp = doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add dataset summary
        doc.add_heading('Dataset Summary', level=1)
        doc.add_paragraph(f"Total Rows: {len(df)}")
        doc.add_paragraph(f"Total Columns: {len(df.columns)}")
        
        # Add results as paragraphs, handling sections properly
        doc.add_heading('Analysis Results', level=1)
        
        current_section = None
        for line in analysis_results:
            # Check if this is a section header
            if line.strip().isupper() or (line.strip() and line.strip()[0] == line.strip()[-1] == "="):
                # Add as a section header (level 2)
                current_section = line.strip().replace("=", "").strip()
                doc.add_heading(current_section, level=2)
            elif line.startswith('---') or not line.strip():
                # Skip separators and blank lines
                continue
            elif line.strip().endswith(':') and len(line.strip()) < 50:
                # Treat as a subsection
                doc.add_heading(line.strip(), level=3)
            else:
                # Add as regular paragraph
                doc.add_paragraph(line)
        
        # Add plots
        if plots and any(plots):
            doc.add_heading('Visualizations', level=1)
            
            # Process each plot type
            for plot_type, plot_data in plots.items():
                if plot_data and len(plot_data) > 0:
                    # Add a heading for this plot type
                    doc.add_heading(plot_type, level=2)
                    
                    for name, img_str in plot_data:
                        doc.add_heading(name, level=3)
                        # Save the image to a temporary file and add it to the document
                        image_data = io.BytesIO(base64.b64decode(img_str))
                        doc.add_picture(image_data, width=Inches(6))
                        doc.add_paragraph()  # Add spacing
        
        # Save the document to a BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes

    # Add Word document download button
    @output
    @render.ui
    def word_download_button():
        return ui.download_button(
            "download_word", 
            "Download Report as Word Document", 
            class_="btn-info mt-2", 
            disabled=not analysis_run()
        )
    
    # Implement Word document download handler
    @output
    @render.download(filename="hfe_analysis_report.docx")
    def download_word():
        # Return None if no analysis has been run
        if not analysis_run():
            return None
        
        # Collect all plots to include in the document
        plots_to_include = {
            "Genotype Distribution": plots_genotype() if plots_genotype() else [],
            "Genotypes by Age": plots_age() if plots_age() else [],
            "Genotypes by Gender": plots_gender() if plots_gender() else [],
            "Genotypes by Diagnosis and Gender": plots_diagnosis_gender() if plots_diagnosis_gender() else [],
            "Risk Distribution": [(None, plot_risk())] if plot_risk() else [],
            "Diagnosis Associations": [(None, plot_diagnosis())] if plot_diagnosis() else [],
            "Liver Disease": [(None, plot_liver())] if plot_liver() else [],
            "Hardy-Weinberg Equilibrium": plots_hardy_weinberg() if plots_hardy_weinberg() else [],
            "Diagnosis Trends": trends_plots_data() if trends_plots_data() else []
        }
        
        # Generate the Word document
        try:
            return generate_word_report(
                analysis_results=analysis_results_text(), 
                df=data(), 
                plots=plots_to_include
            )
        except Exception as e:
            import traceback
            error_traceback = traceback.format_exc()
            ui.notification_show(
                f"Error generating Word document: {str(e)}", 
                type="error", 
                duration=None
            )
            return None

# Create the Shiny app
app = App(create_ui(), server)

# Run the app if this script is executed directly
if __name__ == "__main__":
    print("Starting app on http://127.0.0.1:8095")
    app.run(host="127.0.0.1", port=8095) 